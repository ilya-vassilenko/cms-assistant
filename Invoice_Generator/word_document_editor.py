#!/usr/bin/env python3
"""
Word Document Editor Class
Handles Word document template processing, placeholder replacement, and document manipulation.
"""

import os
import shutil
import subprocess
from datetime import datetime, timedelta
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional

class WordDocumentEditor:
    """
    A class to handle Word document template processing and editing.
    
    Attributes:
        template_path (str): Path to the Word document template
        document (Document): The loaded Word document
    """
    
    def __init__(self, template_path: str):
        """
        Initialize the WordDocumentEditor.
        
        Args:
            template_path (str): Path to the Word document template
        """
        self.template_path = template_path
        self.document = None
        self._validate_template()
    
    def _validate_template(self):
        """Validate that the template file exists and is accessible."""
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template file '{self.template_path}' not found!")
    
    def load_document(self) -> bool:
        """
        Load the Word document template.
        
        Returns:
            bool: True if document loaded successfully, False otherwise
        """
        try:
            self.document = Document(self.template_path)
            return True
        except Exception as e:
            print(f"Error loading document: {e}")
            return False
    
    @staticmethod
    def get_today_formatted() -> str:
        """Get today's date in MMMM dd, YYYY format"""
        return datetime.now().strftime("%B %d, %Y")
    
    @staticmethod
    def get_last_month_formatted() -> str:
        """Get last month in MMMM YYYY format"""
        today = datetime.now()
        # Get first day of current month, then subtract one day to get last month
        first_day_current = today.replace(day=1)
        last_month = first_day_current - timedelta(days=1)
        return last_month.strftime("%B %Y")
    
    @staticmethod
    def get_pay_by_date_formatted() -> str:
        """Get today + 30 days in MMMM dd, YYYY format"""
        today = datetime.now()
        pay_by_date = today + timedelta(days=30)
        return pay_by_date.strftime("%B %d, %Y")
    
    @staticmethod
    def get_last_month_date() -> datetime:
        """Get last month date object for folder naming"""
        today = datetime.now()
        first_day_current = today.replace(day=1)
        return first_day_current - timedelta(days=1)
    
    def find_and_replace_text(self, old_text: str, new_text: str) -> int:
        """
        Find and replace text in all paragraphs and tables.
        
        Args:
            old_text (str): Text to find and replace
            new_text (str): Text to replace with
            
        Returns:
            int: Number of replacements made
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        replacements_made = 0
        
        # Replace in paragraphs
        for paragraph in self.document.paragraphs:
            if old_text in paragraph.text:
                # Clear the paragraph and rebuild it with the replacement
                full_text = paragraph.text
                if old_text in full_text:
                    paragraph.clear()
                    paragraph.add_run(full_text.replace(old_text, new_text))
                    replacements_made += 1
        
        # Replace in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            # Clear the paragraph and rebuild it with the replacement
                            full_text = paragraph.text
                            if old_text in full_text:
                                paragraph.clear()
                                paragraph.add_run(full_text.replace(old_text, new_text))
                                replacements_made += 1
        
        return replacements_made
    
    def replace_date_placeholders(self) -> Dict[str, int]:
        """
        Replace all date placeholders in the document.
        
        Returns:
            Dict[str, int]: Dictionary with placeholder names and replacement counts
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Get formatted dates
        today_formatted = self.get_today_formatted()
        last_month_formatted = self.get_last_month_formatted()
        pay_by_date_formatted = self.get_pay_by_date_formatted()
        
        # Replace placeholders
        today_replacements = self.find_and_replace_text("[TODAY]", today_formatted)
        last_month_replacements = self.find_and_replace_text("[LAST_MONTH]", last_month_formatted)
        pay_by_date_replacements = self.find_and_replace_text("[PAY_BY_DATE]", pay_by_date_formatted)
        
        return {
            "TODAY": today_replacements,
            "LAST_MONTH": last_month_replacements,
            "PAY_BY_DATE": pay_by_date_replacements,
            "total": today_replacements + last_month_replacements + pay_by_date_replacements
        }
    
    def replace_rate_placeholder(self, currency: str, hourly_rate: float) -> int:
        """
        Replace [RATE] placeholder with formatted hourly rate and currency.
        
        Args:
            currency (str): Currency code (e.g., "CHF")
            hourly_rate (float): Hourly rate value
            
        Returns:
            int: Number of replacements made
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Format the rate as "CHF 170" or "USD 150"
        rate_formatted = f"{currency} {hourly_rate:.0f}"
        
        # Replace [RATE] placeholder
        rate_replacements = self.find_and_replace_text("[RATE]", rate_formatted)
        
        return rate_replacements
    
    def add_table_entries(self, table_name: str, entries: List[Dict]) -> bool:
        """
        Add entries to a specific table in the document.
        
        Args:
            table_name (str): Name or identifier of the table to find
            entries (List[Dict]): List of entries to add, each with keys: date, topic, efforts, hours
            
        Returns:
            bool: True if entries were added successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Find the table (simplified - assumes first table for now)
        table = None
        for t in self.document.tables:
            table = t
            break
        
        if not table:
            print(f"Warning: Could not find table '{table_name}'")
            return False
        
        try:
            # Add new entries after the header row (first row)
            for entry in entries:
                new_row = table.add_row()
                new_row.cells[0].text = entry.get('date', '')
                new_row.cells[1].text = entry.get('topic', '')
                new_row.cells[2].text = entry.get('efforts', '')
                new_row.cells[3].text = str(entry.get('hours', 0))
            
            print(f"Added {len(entries)} entries to the table")
            return True
            
        except Exception as e:
            print(f"Error adding table entries: {e}")
            return False
    
    def add_rows_at_bottom(self, num_rows: int) -> bool:
        """
        Find the first table and add X rows at the bottom.
        
        Args:
            num_rows (int): Number of rows to add at the bottom
            
        Returns:
            bool: True if rows were added successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Find the first table
        if not self.document.tables:
            print("Warning: No tables found in the document")
            return False
        
        table = self.document.tables[0]
        
        try:
            total_rows = len(table.rows)
            
            if total_rows == 0:
                print("Warning: Table has no rows")
                return False
            
            print(f"Table has {total_rows} rows, adding {num_rows} rows at the bottom")
            
            # Get the number of columns from the first row
            num_cols = len(table.rows[0].cells)
            
            # Add rows at the bottom
            for i in range(num_rows):
                new_row = table.add_row()
                # Ensure the new row has the same number of cells as other rows
                while len(new_row.cells) < num_cols:
                    new_row._element.append(new_row._element._new_tc())
            
            print(f"Successfully added {num_rows} rows to the table")
            print(f"Table now has {len(table.rows)} rows")
            return True
            
        except Exception as e:
            print(f"Error adding rows to table: {e}")
            return False
    
    def add_working_item_to_first_free_row(self, date: str, topic: str, efforts: str, hours: float) -> bool:
        """
        Add a working item to the first free row from the top in the first table.
        
        Args:
            date (str): Date of the work item
            topic (str): Topic of the work item
            efforts (str): Description of efforts/work done
            hours (float): Number of hours worked
            
        Returns:
            bool: True if item was added successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Find the first table
        if not self.document.tables:
            print("Warning: No tables found in the document")
            return False
        
        table = self.document.tables[0]
        
        try:
            # Find the first free row (starting from row 1, skipping header row 0)
            free_row_index = None
            
            for row_idx in range(1, len(table.rows)):
                # Check if the first cell (date column) is empty
                first_cell_text = table.rows[row_idx].cells[0].text.strip()
                if not first_cell_text:
                    free_row_index = row_idx
                    break
            
            # If no free row found, add a new row at the end
            if free_row_index is None:
                new_row = table.add_row()
                free_row_index = len(table.rows) - 1
            
            # Fill the free row with the working item data
            row = table.rows[free_row_index]
            
            # Ensure the row has enough cells
            while len(row.cells) < 4:
                row._element.append(row._element._new_tc())
            
            # Set the cell values
            row.cells[0].text = str(date)
            row.cells[1].text = str(topic)
            row.cells[2].text = str(efforts)
            row.cells[3].text = str(hours)
            
            print(f"Added working item to row {free_row_index + 1}: {date} - {topic}")
            return True
            
        except Exception as e:
            print(f"Error adding working item to table: {e}")
            return False
    
    def set_last_row_totals(self, total_hours: float) -> bool:
        """
        Set the last row of the first table to show "TOTAL" in the first column
        and the total hours in the fourth column, both in bold font.
        
        Args:
            total_hours (float): Total number of hours to display
            
        Returns:
            bool: True if totals were set successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Find the first table
        if not self.document.tables:
            print("Warning: No tables found in the document")
            return False
        
        table = self.document.tables[0]
        
        try:
            total_rows = len(table.rows)
            if total_rows == 0:
                print("Warning: Table has no rows")
                return False
            
            last_row = table.rows[total_rows - 1]
            
            # Set "TOTAL" in the first column (index 0)
            if len(last_row.cells) > 0:
                # Clear existing content and add "TOTAL" in bold
                first_cell = last_row.cells[0]
                first_cell.text = ""  # Clear existing content
                paragraph = first_cell.paragraphs[0]
                run = paragraph.add_run("TOTAL")
                run.bold = True
            
            # Set total hours in the fourth column (index 3)
            if len(last_row.cells) > 3:
                # Clear existing content and add total hours in bold
                fourth_cell = last_row.cells[3]
                fourth_cell.text = ""  # Clear existing content
                paragraph = fourth_cell.paragraphs[0]
                run = paragraph.add_run(f"{total_hours:.2f}")
                run.bold = True
            
            print(f"Set last row totals: 'TOTAL' in column 1, '{total_hours:.2f}' hours in column 4")
            return True
            
        except Exception as e:
            print(f"Error setting last row totals: {e}")
            return False
    
    def format_table(self) -> bool:
        """
        Format the first table with specified colors, font size, and alignment.
        - Set background color of the last row to #00B050 (green)
        - Set font color of the last row to #FFFFFF (white)
        - Set font size for all rows to 9pt
        - Set text alignment of column 4 (hours) to center
        
        Returns:
            bool: True if formatting was applied successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        # Find the first table
        if not self.document.tables:
            print("Warning: No tables found in the document")
            return False
        
        table = self.document.tables[0]
        
        try:
            total_rows = len(table.rows)
            if total_rows == 0:
                print("Warning: Table has no rows")
                return False
            
            print(f"Formatting table with {total_rows} rows...")
            
            # Define colors
            green_bg = RGBColor(0, 176, 80)  # #00B050
            white_font = RGBColor(255, 255, 255)  # #FFFFFF
            
            # Format all rows
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Set font size to 9pt for all paragraphs in the cell
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                    
                    # Set center alignment for column 4 (hours column, index 3)
                    if cell_idx == 3:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Special formatting for the last row
                    if row_idx == total_rows - 1:
                        # Set background color for the last row
                        cell._tc.get_or_add_tcPr().append(
                            self._create_shading_element(green_bg)
                        )
                        
                        # Set font color to white for the last row
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = white_font
            
            print(f"Successfully formatted table:")
            print(f"- Set font size to 9pt for all rows")
            print(f"- Set center alignment for column 4 (hours)")
            print(f"- Set last row background to green (#00B050)")
            print(f"- Set last row font color to white (#FFFFFF)")
            return True
            
        except Exception as e:
            print(f"Error formatting table: {e}")
            return False
    
    def format_payment_instruction(self) -> bool:
        """
        Find and make bold the payment instruction line containing
        "Please transfer ... to the following bank account by at latest ...:"
        
        Returns:
            bool: True if formatting was applied successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        try:
            payment_instruction_found = False
            
            # Search through all paragraphs for the payment instruction
            for paragraph in self.document.paragraphs:
                paragraph_text = paragraph.text.lower()
                
                # Check if this paragraph contains payment instruction keywords
                if ("please transfer" in paragraph_text and 
                    "bank account" in paragraph_text and 
                    "at latest" in paragraph_text):
                    
                    # Make all runs in this paragraph bold
                    for run in paragraph.runs:
                        run.bold = True
                    
                    print(f"Made payment instruction bold: '{paragraph.text[:50]}...'")
                    payment_instruction_found = True
                    break
            
            if not payment_instruction_found:
                print("Warning: Payment instruction line not found")
                print("Looking for text containing 'Please transfer', 'bank account', and 'at latest'")
                return False
            
            return True
            
        except Exception as e:
            print(f"Error formatting payment instruction: {e}")
            return False
    
    def _create_shading_element(self, color: RGBColor):
        """
        Create a shading element for table cell background color.
        
        Args:
            color (RGBColor): The background color to apply
            
        Returns:
            The shading element
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Convert RGBColor to hex string
        # RGBColor stores values as integers, access them directly
        hex_color = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
        
        # Create the shading element
        shading_xml = f'''
        <w:shd {nsdecls('w')} w:fill="{hex_color}" />
        '''
        return parse_xml(shading_xml)
    
    def save_document(self, output_path: str) -> bool:
        """
        Save the document to the specified path.
        
        Args:
            output_path (str): Path where to save the document
            
        Returns:
            bool: True if document saved successfully, False otherwise
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        try:
            # Create output directory if it doesn't exist
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # Validate document before saving
            if not self._validate_document():
                print("Warning: Document validation failed, but attempting to save anyway...")
            
            # Save the document
            self.document.save(output_path)
            
            # Set proper file permissions (readable and writable by owner)
            import stat
            os.chmod(output_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IROTH)
            
            # Verify the file was created and is readable
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"Document saved successfully: {output_path}")
                print(f"File size: {os.path.getsize(output_path)} bytes")
                return True
            else:
                print(f"Error: Document was not saved properly or file is empty")
                return False
                
        except Exception as e:
            print(f"Error saving document: {e}")
            print(f"Output path: {output_path}")
            print(f"Directory exists: {os.path.exists(os.path.dirname(output_path))}")
            return False
    
    def convert_to_pdf(self, word_path: str) -> bool:
        """
        Convert a Word document to PDF using LibreOffice.
        
        Args:
            word_path (str): Path to the Word document to convert
            
        Returns:
            bool: True if conversion was successful, False otherwise
        """
        if not os.path.exists(word_path):
            print(f"Error: Word document not found: {word_path}")
            return False
        
        try:
            # Generate PDF path (replace .docx with .pdf)
            pdf_path = word_path.replace('.docx', '.pdf')
            
            # Get the directory of the Word document
            output_dir = os.path.dirname(word_path)
            
            print(f"Converting Word document to PDF...")
            print(f"Source: {word_path}")
            print(f"Target: {pdf_path}")
            
            # Use LibreOffice to convert Word to PDF
            # --headless: Run without GUI
            # --convert-to pdf: Convert to PDF format
            # --outdir: Specify output directory
            
            # Try different LibreOffice command paths
            libreoffice_commands = [
                'libreoffice',  # Standard command
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS app bundle
                '/usr/bin/libreoffice',  # Linux/Ubuntu
                '/usr/local/bin/libreoffice'  # Homebrew on Intel Mac
            ]
            
            cmd = None
            for libreoffice_cmd in libreoffice_commands:
                try:
                    # Test if the command exists
                    subprocess.run([libreoffice_cmd, '--version'], 
                                 capture_output=True, timeout=5)
                    cmd = [
                        libreoffice_cmd,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', output_dir,
                        word_path
                    ]
                    print(f"Using LibreOffice command: {libreoffice_cmd}")
                    break
                except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError):
                    continue
            
            if cmd is None:
                raise FileNotFoundError("LibreOffice command not found")
            
            # Run the conversion command
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                # Verify the PDF was created
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    print(f"PDF created successfully: {pdf_path}")
                    print(f"PDF file size: {os.path.getsize(pdf_path)} bytes")
                    return True
                else:
                    print(f"Error: PDF file was not created or is empty")
                    return False
            else:
                print(f"Error: LibreOffice conversion failed")
                print(f"Return code: {result.returncode}")
                print(f"Error output: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            print("Error: PDF conversion timed out after 60 seconds")
            return False
        except FileNotFoundError:
            print("Error: LibreOffice not found. Please install LibreOffice to convert documents to PDF.")
            print("Installation instructions:")
            print("- macOS: brew install --cask libreoffice")
            print("- Ubuntu/Debian: sudo apt-get install libreoffice")
            print("- Windows: Download from https://www.libreoffice.org/")
            return False
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            return False
    
    def copy_pdf_to_folder(self, pdf_path: str, copy_to_folder: str, output_folder_name: str) -> bool:
        """
        Copy the PDF to a specified folder, creating a subfolder with the output folder name.
        
        Args:
            pdf_path (str): Path to the PDF file to copy
            copy_to_folder (str): Base folder where to copy the PDF
            output_folder_name (str): Name of the subfolder to create (same as output folder name)
            
        Returns:
            bool: True if copy was successful, False otherwise
        """
        if not os.path.exists(pdf_path):
            print(f"Error: PDF file not found: {pdf_path}")
            return False
        
        if not copy_to_folder:
            print("Warning: No copy folder specified, skipping PDF copy")
            return True
        
        try:
            # Create the destination subfolder
            destination_subfolder = os.path.join(copy_to_folder, output_folder_name)
            
            print(f"Creating destination folder: {destination_subfolder}")
            os.makedirs(destination_subfolder, exist_ok=True)
            
            # Get the PDF filename
            pdf_filename = os.path.basename(pdf_path)
            destination_path = os.path.join(destination_subfolder, pdf_filename)
            
            print(f"Copying PDF...")
            print(f"Source: {pdf_path}")
            print(f"Destination: {destination_path}")
            
            # Copy the PDF file
            shutil.copy2(pdf_path, destination_path)
            
            # Verify the copy was successful
            if os.path.exists(destination_path) and os.path.getsize(destination_path) > 0:
                print(f"PDF copied successfully!")
                print(f"Copied file size: {os.path.getsize(destination_path)} bytes")
                return True
            else:
                print(f"Error: PDF copy verification failed")
                return False
                
        except Exception as e:
            print(f"Error copying PDF: {e}")
            return False
    
    def get_template_filename(self) -> str:
        """
        Get the template filename.
        
        Returns:
            str: The template filename
        """
        return os.path.basename(self.template_path)
    
    def generate_output_filename(self, last_month_formatted: str) -> str:
        """
        Generate output filename by replacing [LAST_MONTH] and [TODAY] in template filename.
        
        Args:
            last_month_formatted (str): Formatted last month string
            
        Returns:
            str: Generated output filename
        """
        template_filename = self.get_template_filename()
        
        # Replace [LAST_MONTH] with the formatted last month
        output_filename = template_filename.replace("[LAST_MONTH]", last_month_formatted)
        
        # Replace [TODAY] with current date in YYYY-mm-dd format
        today_formatted = datetime.now().strftime('%Y-%m-%d')
        output_filename = output_filename.replace("[TODAY]", today_formatted)
        
        return output_filename
    
    def get_document_info(self) -> Dict[str, str]:
        """
        Get information about the loaded document.
        
        Returns:
            Dict[str, str]: Document information
        """
        if not self.document:
            return {"status": "not_loaded"}
        
        return {
            "status": "loaded",
            "template_path": self.template_path,
            "paragraphs_count": len(self.document.paragraphs),
            "tables_count": len(self.document.tables)
        }
    
    def read_table_cell(self, row: int, col: int) -> str:
        """
        Read a specific cell from the first table.
        
        Args:
            row (int): Row index (0-based)
            col (int): Column index (0-based, where A=0, B=1, C=2, D=3, E=4)
            
        Returns:
            str: Cell content as string
            
        Raises:
            ValueError: If table doesn't exist or cell is out of bounds
        """
        if not self.document:
            raise ValueError("Document not loaded. Call load_document() first.")
        
        if not self.document.tables:
            raise ValueError("No tables found in the document")
        
        table = self.document.tables[0]
        
        if row >= len(table.rows):
            raise ValueError(f"Row {row} is out of bounds. Table has {len(table.rows)} rows.")
        
        row_obj = table.rows[row]
        
        if col >= len(row_obj.cells):
            raise ValueError(f"Column {col} is out of bounds. Row has {len(row_obj.cells)} columns.")
        
        return row_obj.cells[col].text.strip()
    
    def _validate_document(self) -> bool:
        """
        Validate the document structure before saving.
        
        Returns:
            bool: True if document is valid, False otherwise
        """
        if not self.document:
            return False
        
        try:
            # Check if document has content
            if len(self.document.paragraphs) == 0 and len(self.document.tables) == 0:
                print("Warning: Document appears to be empty")
                return False
            
            # Check for any obvious corruption indicators
            for paragraph in self.document.paragraphs:
                if paragraph.text and len(paragraph.text) > 10000:  # Very long paragraphs might indicate corruption
                    print("Warning: Found unusually long paragraph, possible corruption")
                    return False
            
            return True
            
        except Exception as e:
            print(f"Document validation error: {e}")
            return False
