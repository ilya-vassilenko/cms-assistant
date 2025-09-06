#!/usr/bin/env python3
"""
Invoice Generator Script
Processes invoice template Word documents by replacing placeholders with current dates.
"""

import os
import sys
import json
import argparse
from datetime import datetime, timedelta
from docx import Document

def get_today_formatted():
    """Get today's date in MMMM dd, YYYY format"""
    return datetime.now().strftime("%B %d, %Y")

def get_last_month_formatted():
    """Get last month in MMMM YYYY format"""
    today = datetime.now()
    # Get first day of current month, then subtract one day to get last month
    first_day_current = today.replace(day=1)
    last_month = first_day_current - timedelta(days=1)
    return last_month.strftime("%B %Y")

def get_pay_by_date_formatted():
    """Get today + 30 days in MMMM dd, YYYY format"""
    today = datetime.now()
    pay_by_date = today + timedelta(days=30)
    return pay_by_date.strftime("%B %d, %Y")

def get_last_month_date():
    """Get last month date object for folder naming"""
    today = datetime.now()
    first_day_current = today.replace(day=1)
    return first_day_current - timedelta(days=1)

def find_and_replace_text(doc, old_text, new_text):
    """Find and replace text in all paragraphs and tables"""
    replacements_made = 0
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            # Clear the paragraph and rebuild it with the replacement
            full_text = paragraph.text
            if old_text in full_text:
                paragraph.clear()
                paragraph.add_run(full_text.replace(old_text, new_text))
                replacements_made += 1
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        # Clear the paragraph and rebuild it with the replacement
                        full_text = paragraph.text
                        if old_text in full_text:
                            paragraph.clear()
                            paragraph.add_run(full_text.replace(old_text, new_text))
                            replacements_made += 1
    
    return replacements_made

def create_invoice_folder(invoice_folder_base, last_month_date):
    """Create folder with format YYYY-MM-DD <last month name> <last month year>"""
    folder_name = f"{last_month_date.strftime('%Y-%m-%d')} {last_month_date.strftime('%B %Y')}"
    folder_path = os.path.join(invoice_folder_base, folder_name)
    
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Created folder: {folder_path}")
    else:
        print(f"Folder already exists: {folder_path}")
    
    return folder_path

def load_config(config_path):
    """Load configuration from JSON file"""
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
        return config
    except FileNotFoundError:
        print(f"Error: Config file '{config_path}' not found!")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in config file: {e}")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description='Generate invoice from template')
    parser.add_argument('--config', help='Path to configuration JSON file')
    args = parser.parse_args()
    
    # Load configuration
    config = load_config(args.config)
    
    # Get template path from config
    template_path = config.get('template')
    if not template_path:
        print("Error: 'template' not found in config file!")
        sys.exit(1)
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"Error: Template file '{template_path}' not found!")
        sys.exit(1)
    
    # Get invoice folder from config
    invoice_folder_base = config.get('invoice_folder')
    if not invoice_folder_base:
        print("Error: 'invoice_folder' not found in config file!")
        sys.exit(1)
    
    try:
        # Load the document
        print("Loading template document...")
        doc = Document(template_path)
        
        # Get formatted dates
        today_formatted = get_today_formatted()
        last_month_formatted = get_last_month_formatted()
        pay_by_date_formatted = get_pay_by_date_formatted()
        last_month_date = get_last_month_date()
        
        # Replace placeholders
        print("Replacing placeholders...")
        today_replacements = find_and_replace_text(doc, "[TODAY]", today_formatted)
        last_month_replacements = find_and_replace_text(doc, "[LAST_MONTH]", last_month_formatted)
        pay_by_date_replacements = find_and_replace_text(doc, "[PAY_BY_DATE]", pay_by_date_formatted)
        
        print(f"Replaced [TODAY] with: {today_formatted} ({today_replacements} replacements)")
        print(f"Replaced [LAST_MONTH] with: {last_month_formatted} ({last_month_replacements} replacements)")
        print(f"Replaced [PAY_BY_DATE] with: {pay_by_date_formatted} ({pay_by_date_replacements} replacements)")
        
        total_replacements = today_replacements + last_month_replacements + pay_by_date_replacements
        if total_replacements == 0:
            print("WARNING: No placeholders were found and replaced in the document!")
            print("Please check that the template contains [TODAY], [LAST_MONTH], and/or [PAY_BY_DATE] placeholders.")
        
        # Create invoice folder
        invoice_folder = create_invoice_folder(invoice_folder_base, last_month_date)
        
        # Generate output filename by replacing [LAST_MONTH] in template filename
        template_filename = os.path.basename(template_path)
        output_filename = template_filename.replace("[LAST_MONTH]", last_month_formatted)
        output_path = os.path.join(invoice_folder, output_filename)
        
        # Save the document
        print(f"Saving document as '{output_path}'...")
        doc.save(output_path)
        print("Invoice generated successfully!")
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
