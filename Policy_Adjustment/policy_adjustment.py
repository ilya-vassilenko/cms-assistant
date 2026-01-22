#!/usr/bin/env python3
"""
Policy Adjustment Script

Processes Word documents by:
- loading config from JSON via --config
- applying search/replace patterns
- decreasing heading levels by one (Heading 1 -> Heading 2, etc.)
- saving results to output folder
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
from email import policy
from email.parser import BytesParser
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

try:
    from colorama import Fore, Style, init as colorama_init
except Exception:  # pragma: no cover
    # Fallback when colorama isn't installed; keep output readable without colors.
    class _NoColor:
        BLACK = RED = GREEN = YELLOW = BLUE = MAGENTA = CYAN = WHITE = RESET = ""

    class _NoStyle:
        RESET_ALL = ""

    Fore = _NoColor()  # type: ignore
    Style = _NoStyle()  # type: ignore

    def colorama_init(*_args, **_kwargs) -> None:  # type: ignore
        return None
from docx import Document


@dataclass(frozen=True)
class Replacement:
    search: str
    replace: str


def c_info(msg: str) -> str:
    return f"{Fore.CYAN}{msg}{Style.RESET_ALL}"


def c_ok(msg: str) -> str:
    return f"{Fore.GREEN}{msg}{Style.RESET_ALL}"


def c_warn(msg: str) -> str:
    return f"{Fore.YELLOW}{msg}{Style.RESET_ALL}"


def c_err(msg: str) -> str:
    return f"{Fore.RED}{msg}{Style.RESET_ALL}"


def load_config(config_path: str) -> dict:
    """Load configuration from JSON file (mirrors Invoice_Generator style)."""
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        print(c_err(f"Error: Config file '{config_path}' not found!"))
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(c_err(f"Error: Invalid JSON in config file: {e}"))
        sys.exit(1)


def parse_replacements(config: dict) -> List[Replacement]:
    reps = []
    raw = config.get("replacements", [])
    if raw is None:
        return reps
    if not isinstance(raw, list):
        raise ValueError("'replacements' must be a list")
    for idx, item in enumerate(raw):
        if not isinstance(item, dict):
            raise ValueError(f"replacements[{idx}] must be an object")
        if "search" not in item or "replace" not in item:
            raise ValueError(f"replacements[{idx}] must contain 'search' and 'replace'")
        reps.append(Replacement(search=str(item["search"]), replace=str(item["replace"])))
    return reps


def iter_paragraphs_in_document(doc: Document):
    # Main body
    for p in doc.paragraphs:
        yield p
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_in_paragraph(paragraph, search: str, replace: str) -> Tuple[int, bool]:
    """
    Replace occurrences in a paragraph.

    Returns (replacements_made, formatting_preserved).

    We first try run-by-run replacement (preserves formatting when search doesn't span runs).
    If the text is found in paragraph.text but not replaced via runs (likely spanning runs),
    we fall back to clearing and re-adding a single run (formatting lost for that paragraph).
    """
    if not search:
        return 0, True

    original_text = paragraph.text or ""
    if search not in original_text:
        return 0, True

    # Attempt run-level replacement
    made = 0
    for run in paragraph.runs:
        if search in run.text:
            run.text = run.text.replace(search, replace)
            made += 1

    # If still present, fallback to paragraph-level rewrite (may lose formatting)
    if search in paragraph.text:
        try:
            new_text = (paragraph.text or "").replace(search, replace)
            paragraph.clear()
            paragraph.add_run(new_text)
            # Count as 1 replacement action (consistent with WordDocumentEditor behavior)
            made += 1
            return made, False
        except Exception:
            # If rewriting fails, keep best-effort changes
            return made, True

    return made, True


def apply_replacements(doc: Document, replacements: List[Replacement]) -> Tuple[int, int]:
    """
    Apply all replacements in document paragraphs and table cell paragraphs.
    Returns: (total_replacement_actions, paragraphs_with_formatting_loss)
    """
    total = 0
    formatting_loss_paragraphs = 0

    for p in iter_paragraphs_in_document(doc):
        for rep in replacements:
            made, preserved = replace_in_paragraph(p, rep.search, rep.replace)
            if made:
                total += made
                if not preserved:
                    formatting_loss_paragraphs += 1

    return total, formatting_loss_paragraphs


def decrease_heading_levels(doc: Document) -> Tuple[int, int]:
    """
    Decrease heading levels by one:
    - Heading 1 -> Heading 2
    - ...
    - Heading 8 -> Heading 9
    Heading 9 has no "Heading 10" in Word; it is left unchanged and counted as skipped.

    Returns: (changed_count, skipped_heading9_count)
    """
    changed = 0
    skipped = 0

    for p in doc.paragraphs:
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            continue

        if not style_name.startswith("Heading "):
            continue

        parts = style_name.split()
        if len(parts) != 2:
            continue

        try:
            level = int(parts[1])
        except ValueError:
            continue

        if level >= 9:
            skipped += 1
            continue

        new_style = f"Heading {level + 1}"
        try:
            p.style = new_style
            changed += 1
        except Exception:
            # Style might not exist in some templates; skip silently.
            continue

    return changed, skipped


def find_word_files(input_folder: Path) -> List[Path]:
    allowed = {".docx", ".doc"}
    files: List[Path] = []
    for root, _, filenames in os.walk(input_folder):
        for fn in filenames:
            p = Path(root) / fn
            if p.suffix.lower() in allowed and not p.name.startswith("~$"):
                files.append(p)
    files.sort()
    return files


def looks_like_zip(path: Path) -> bool:
    try:
        with open(path, "rb") as f:
            return f.read(2) == b"PK"
    except Exception:
        return False


def looks_like_ole_compound_doc(path: Path) -> bool:
    """
    Legacy Word .doc (binary) is an OLE compound file with magic:
    D0 CF 11 E0 A1 B1 1A E1
    """
    try:
        with open(path, "rb") as f:
            return f.read(8) == bytes.fromhex("D0CF11E0A1B11AE1")
    except Exception:
        return False


def try_extract_html_from_mime_message(path: Path) -> Optional[str]:
    """
    Some ".doc" inputs are actually EML/MIME messages (e.g., Confluence exports)
    containing a text/html part. If so, return the decoded HTML string.
    """
    try:
        raw = path.read_bytes()
        msg = BytesParser(policy=policy.default).parsebytes(raw)
        if not msg.is_multipart():
            return None
        for part in msg.walk():
            if part.get_content_type() == "text/html":
                html = part.get_content()  # decodes charset + quoted-printable/base64
                if isinstance(html, str) and "<html" in html.lower():
                    return html
        return None
    except Exception:
        return None


def _print_conversion_debug(debug_lines: List[str], verbose: bool) -> None:
    if not verbose:
        return
    print(c_info("LibreOffice conversion debug:"))
    for line in debug_lines:
        print(c_info(f"- {line}"))


def try_convert_to_docx(source_file: Path, tmpdir_path: Path, *, verbose: bool = False) -> Optional[Path]:
    """
    Best-effort conversion to .docx using LibreOffice (soffice/libreoffice).
    If unavailable or conversion fails, returns None.
    """
    debug: List[str] = []
    tmp_in = tmpdir_path / source_file.name
    try:
        if source_file.resolve() != tmp_in.resolve():
            shutil.copy2(source_file, tmp_in)
    except Exception:
        # Best-effort: if we can't resolve/copy, still attempt to convert using the original path.
        tmp_in = source_file
    debug.append(f"source_file={source_file}")
    debug.append(f"tmp_in={tmp_in}")
    debug.append(f"tmpdir_path={tmpdir_path}")

    candidates = [
        "libreoffice",
        "soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/libreoffice",
        "/usr/local/bin/libreoffice",
    ]

    soffice_cmd = None
    for c in candidates:
        try:
            ver = subprocess.run([c, "--version"], capture_output=True, text=True, timeout=5, check=False)
            debug.append(f"candidate={c} version_rc={ver.returncode}")
            if ver.stdout:
                debug.append(f"candidate={c} version_stdout={ver.stdout.strip()[:500]}")
            if ver.stderr:
                debug.append(f"candidate={c} version_stderr={ver.stderr.strip()[:500]}")
            if ver.returncode == 0:
                soffice_cmd = c
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            debug.append(f"candidate={c} not_found_or_timeout")
            continue

    if soffice_cmd is None:
        _print_conversion_debug(debug + ["No working LibreOffice/soffice command found."], verbose)
        return None

    try:
        # LibreOffice sometimes needs an explicit export filter name for docx.
        # We'll try a small set of known-good filter specs.
        convert_specs = [
            "docx",
            'docx:"MS Word 2007 XML"',
            'docx:"Office Open XML Text"',
        ]

        last_stdout = ""
        last_stderr = ""
        last_rc = -1

        for spec in convert_specs:
            produced = tmpdir_path / (source_file.stem + ".docx")
            if produced.exists():
                try:
                    produced.unlink()
                except Exception:
                    pass

            cmd = [soffice_cmd, "--headless", "--convert-to", spec, "--outdir", str(tmpdir_path), str(tmp_in)]
            debug.append(f"convert_cmd={' '.join(str(x) for x in cmd)}")

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, check=False)
            last_rc = result.returncode
            last_stdout = (result.stdout or "").strip()
            last_stderr = (result.stderr or "").strip()

            debug.append(f"convert_rc={last_rc}")
            if last_stdout:
                debug.append(f"convert_stdout={last_stdout[:2000]}")
            if last_stderr:
                debug.append(f"convert_stderr={last_stderr[:2000]}")

            # Some LO builds return rc=0 but still print an error and produce nothing.
            if last_rc != 0 or "no export filter" in last_stderr.lower():
                continue

            debug.append(f"expected_produced={produced}")
            if produced.exists() and produced.stat().st_size > 0:
                return produced

        # Failed all conversion attempts; dump debug.
        try:
            listing = ", ".join(sorted(p.name for p in tmpdir_path.iterdir())[:50])
        except Exception:
            listing = "<failed to list tmpdir>"
        debug.append(f"tmpdir_listing={listing}")
        _print_conversion_debug(debug, verbose)
        return None
    except subprocess.TimeoutExpired:
        _print_conversion_debug(debug + ["convert_error=timeout"], verbose)
        return None
    except FileNotFoundError:
        _print_conversion_debug(debug + ["convert_error=command_not_found"], verbose)
        return None


def process_one_file(source_path: Path, output_folder: Path, replacements: List[Replacement], *, verbose: bool = False) -> bool:
    print(c_info(f"\nProcessing: {source_path}"))

    load_path = source_path
    temp_dir_ctx = None

    if source_path.suffix.lower() == ".doc":
        # A surprising number of ".doc" files in the wild are not actually Word binary docs.
        # Handle common cases explicitly to avoid "destroyed" outputs.
        if looks_like_zip(source_path):
            print(c_warn("File has .doc extension but looks like a ZIP (likely misnamed .docx). Loading as .docx..."))
        else:
            html = try_extract_html_from_mime_message(source_path)
            if html is not None:
                print(c_warn("'.doc' looks like a MIME/HTML export (e.g., Confluence). Extracting HTML and converting to .docx..."))
                temp_dir_ctx = tempfile.TemporaryDirectory()
                tmpdir_path = Path(temp_dir_ctx.name)
                html_path = tmpdir_path / f"{source_path.stem}.html"
                html_path.write_text(html, encoding="utf-8")
                converted = try_convert_to_docx(html_path, tmpdir_path, verbose=verbose)
                if converted is None:
                    print(c_err("Failed to convert extracted HTML -> .docx (LibreOffice not found or conversion failed). Skipping."))
                    print(c_info("Tip: install LibreOffice to enable conversion (e.g. `brew install --cask libreoffice`)."))
                    temp_dir_ctx.cleanup()
                    return False
                load_path = converted
                print(c_ok(f"Converted to: {load_path}"))
            else:
                # Some exports are just HTML saved with a .doc extension (not MIME-wrapped).
                try:
                    head = source_path.read_text(encoding="utf-8", errors="ignore")[:4096].lower()
                except Exception:
                    head = ""

                if "<html" in head:
                    print(c_warn("'.doc' appears to be HTML content. Converting HTML -> .docx via LibreOffice..."))
                    temp_dir_ctx = tempfile.TemporaryDirectory()
                    tmpdir_path = Path(temp_dir_ctx.name)
                    converted = try_convert_to_docx(source_path, tmpdir_path, verbose=verbose)
                    if converted is None:
                        print(c_err("Failed to convert HTML -> .docx (LibreOffice not found or conversion failed). Skipping."))
                        print(c_info("Tip: install LibreOffice to enable conversion (e.g. `brew install --cask libreoffice`)."))
                        temp_dir_ctx.cleanup()
                        return False
                    load_path = converted
                    print(c_ok(f"Converted to: {load_path}"))
                elif looks_like_ole_compound_doc(source_path):
                    print(c_warn("Legacy Word binary .doc detected; converting to .docx via LibreOffice..."))
                    temp_dir_ctx = tempfile.TemporaryDirectory()
                    tmpdir_path = Path(temp_dir_ctx.name)
                    converted = try_convert_to_docx(source_path, tmpdir_path, verbose=verbose)
                    if converted is None:
                        print(c_err("Failed to convert .doc -> .docx (LibreOffice not found or conversion failed). Skipping."))
                        print(c_info("Tip: install LibreOffice to enable .doc support (e.g. `brew install --cask libreoffice`)."))
                        temp_dir_ctx.cleanup()
                        return False
                    load_path = converted
                    print(c_ok(f"Converted to: {load_path}"))
                else:
                    print(c_err("'.doc' does not look like a Word binary .doc (OLE) and is not a recognized MIME/HTML export. Skipping."))
                    print(c_info("If this is HTML/RTF, convert it to .docx first or export as .docx."))
                    return False

    try:
        doc = Document(str(load_path))
    except Exception as e:
        print(c_err(f"Failed to load document: {e}"))
        if temp_dir_ctx is not None:
            temp_dir_ctx.cleanup()
        return False

    rep_total, formatting_loss = apply_replacements(doc, replacements)
    headings_changed, headings_skipped = decrease_heading_levels(doc)

    output_folder.mkdir(parents=True, exist_ok=True)
    out_name = source_path.name
    if source_path.suffix.lower() == ".doc":
        out_name = f"{source_path.stem}.docx"
    out_path = output_folder / out_name

    try:
        doc.save(str(out_path))
    except Exception as e:
        print(c_err(f"Failed to save output: {e}"))
        if temp_dir_ctx is not None:
            temp_dir_ctx.cleanup()
        return False
    finally:
        if temp_dir_ctx is not None:
            temp_dir_ctx.cleanup()

    print(c_ok(f"Saved: {out_path}"))
    print(c_info(f"Replacements actions: {rep_total}"))
    if formatting_loss:
        print(c_warn(f"Paragraphs rewritten (formatting may be lost): {formatting_loss}"))
    print(c_info(f"Heading level changes: {headings_changed}"))
    if headings_skipped:
        print(c_warn(f"Heading 9 skipped (no Heading 10): {headings_skipped}"))

    return True


def main() -> None:
    colorama_init()

    parser = argparse.ArgumentParser(description="Adjust policies in Word documents")
    parser.add_argument("--config", required=True, help="Path to configuration JSON file")
    parser.add_argument("-v", "--verbose", action="store_true", help="Print verbose debug details (e.g., conversion logs)")
    args = parser.parse_args()
    # If you're running under VS Code / debugpy, always enable verbose so errors show in Debug Console.
    verbose = bool(args.verbose or sys.gettrace())

    config_path = Path(args.config)
    config = load_config(str(config_path))

    input_folder = config.get("input_folder")
    output_folder = config.get("output_folder")
    if not input_folder or not output_folder:
        print(c_err("Error: 'input_folder' and 'output_folder' must be set in the config file!"))
        sys.exit(1)

    try:
        replacements = parse_replacements(config)
    except ValueError as e:
        print(c_err(f"Error in config: {e}"))
        sys.exit(1)

    workspace_root = Path(__file__).resolve().parents[1]

    def resolve_config_path(p: str) -> Path:
        """
        Resolve paths in a forgiving way:
        - absolute paths stay absolute
        - relative paths are tried relative to config file dir, then relative to repo root
        """
        candidate = Path(p)
        if candidate.is_absolute():
            return candidate

        via_config = (config_path.parent / candidate).resolve()
        if via_config.exists():
            return via_config

        via_repo = (workspace_root / candidate).resolve()
        return via_repo

    input_path = resolve_config_path(str(input_folder))
    output_path = resolve_config_path(str(output_folder))

    if not input_path.exists() or not input_path.is_dir():
        print(c_err(f"Error: input_folder does not exist or is not a directory: {input_path}"))
        sys.exit(1)

    files = find_word_files(input_path)
    if not files:
        print(c_warn(f"No .docx/.doc files found under: {input_path}"))
        sys.exit(0)

    print(c_info(f"Found {len(files)} file(s) to process."))
    if replacements:
        print(c_info(f"Loaded {len(replacements)} replacement rule(s)."))
    else:
        print(c_warn("No replacements configured (replacements=[]). Only heading adjustment will run."))

    ok = 0
    fail = 0
    for f in files:
        if process_one_file(f, output_path, replacements, verbose=verbose):
            ok += 1
        else:
            fail += 1

    print(c_info("\n" + "=" * 60))
    print(c_ok(f"Done. Success: {ok}"))
    if fail:
        print(c_err(f"Failed: {fail}"))
        sys.exit(2)


if __name__ == "__main__":
    main()

